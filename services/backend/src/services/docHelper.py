#!/usr/bin/env python3
# @pattern Template Method
# -*- coding: utf-8 -*-
"""
KHY OS Document Helper — PDF-to-Word conversion and OCR text extraction.
Called by Node.js tools via subprocess. Outputs JSON to stdout.
Dependencies: pip install khy-os[doc]
"""
import sys
import json
import os
import shutil
import subprocess
import tempfile

MAX_PDF_SIZE = 50 * 1024 * 1024  # 50 MB
MAX_IMAGE_SIZE = 20 * 1024 * 1024  # 20 MB


def _output(data):
    """Print JSON result to stdout for Node.js to parse."""
    print(json.dumps(data, ensure_ascii=False))


def check_deps():
    """Check which optional dependencies are available."""
    status = {
        "pdf2docx": False,
        "pytesseract": False,
        "pillow": False,
        "tesseract_binary": False,
    }
    try:
        import pdf2docx  # noqa: F401
        status["pdf2docx"] = True
    except ImportError:
        pass
    try:
        from PIL import Image  # noqa: F401
        status["pillow"] = True
    except ImportError:
        pass
    try:
        import pytesseract
        status["pytesseract"] = True
        # Verify tesseract binary is installed
        pytesseract.get_tesseract_version()
        status["tesseract_binary"] = True
    except ImportError:
        pass
    except Exception:
        # pytesseract installed but tesseract binary missing
        status["tesseract_binary"] = False

    _output({"success": True, "deps": status})


def pdf_to_word(input_path, output_path):
    """Convert PDF to Word (.docx) using pdf2docx."""
    try:
        from pdf2docx import Converter
    except ImportError:
        _output({
            "success": False,
            "error": "pdf2docx library not installed. Run: pip install khy-os[doc]",
        })
        return

    if not os.path.isfile(input_path):
        _output({"success": False, "error": f"File not found: {input_path}"})
        return

    size = os.path.getsize(input_path)
    if size > MAX_PDF_SIZE:
        _output({
            "success": False,
            "error": f"PDF too large: {size / (1024 * 1024):.1f}MB (max 50MB)",
        })
        return

    try:
        # Ensure output directory exists
        out_dir = os.path.dirname(output_path)
        if out_dir and not os.path.isdir(out_dir):
            os.makedirs(out_dir, exist_ok=True)

        cv = Converter(input_path)
        cv.convert(output_path)
        cv.close()

        out_size = os.path.getsize(output_path)
        _output({
            "success": True,
            "output": output_path,
            "inputSize": size,
            "outputSize": out_size,
            "message": f"Converted: {os.path.basename(output_path)} ({out_size / 1024:.0f}KB)",
        })
    except Exception as e:
        _output({"success": False, "error": f"PDF conversion failed: {str(e)}"})


def _tesseract_available_langs(bin_path):
    """Return the trained languages the tesseract binary actually has."""
    try:
        proc = subprocess.run(
            [bin_path, "--list-langs"],
            capture_output=True, text=True, timeout=10,
        )
    except Exception:
        return []
    langs = []
    for line in (proc.stdout or "").splitlines():
        line = line.strip()
        if not line or line.lower().startswith("list of available"):
            continue
        langs.append(line)
    return langs


def _resolve_lang(requested, available):
    """Narrow a requested 'a+b+c' lang spec to languages that are installed.

    Tesseract fails to load a language whose traineddata is missing (e.g.
    'chi_sim' when only 'eng' is present), so we keep only the usable parts and
    fall back to English (then any non-osd language) rather than erroring out.
    """
    parts = [p for p in str(requested or "").split("+") if p]
    avail = set(available or [])
    if not avail:
        # Cannot introspect — trust the caller's request as-is.
        return requested or "eng"
    usable = [p for p in parts if p in avail]
    if usable:
        return "+".join(usable)
    if "eng" in avail:
        return "eng"
    for lang in available or []:
        if lang != "osd":
            return lang
    return "eng"


def _mean_tsv_confidence(tsv_path):
    """Average the positive per-word confidences from a tesseract .tsv file.

    tesseract's TSV output has a ``conf`` column (0-100 per recognized word, -1
    for layout/non-word rows). We average only the strictly-positive values, the
    same shape the pytesseract path uses. Returns a float mean, or ``None`` when
    the file is missing / has no ``conf`` column / has no usable rows — so the
    caller can honestly decline to fabricate a score.
    """
    try:
        if not os.path.isfile(tsv_path):
            return None
        confs = []
        with open(tsv_path, "r", encoding="utf-8", errors="replace") as fh:
            header = fh.readline().rstrip("\n").split("\t")
            try:
                conf_idx = header.index("conf")
            except ValueError:
                return None
            for line in fh:
                cols = line.rstrip("\n").split("\t")
                if len(cols) <= conf_idx:
                    continue
                raw = cols[conf_idx].strip()
                if not raw.lstrip("-").isdigit():
                    continue
                value = int(raw)
                if value > 0:
                    confs.append(value)
        if not confs:
            return None
        return sum(confs) / len(confs)
    except Exception:  # noqa: BLE001 — confidence is best-effort, never fatal
        return None


def _ocr_via_cli_with_confidence(bin_path, image_path, eff_lang):
    """Single tesseract pass that emits both text (.txt) and per-word conf (.tsv).

    ``tesseract <img> <outbase> -l <lang> txt tsv`` writes outbase.txt (the exact
    same rendered text as the ``stdout`` txt pass) AND outbase.tsv (per-word
    confidence) in one invocation. This lets the CLI fallback report an honest
    averaged confidence — and therefore a truthful ``needsAiFallback`` — just like
    the pytesseract path, instead of a hardcoded ``confidence: 0``.

    Returns a result dict on success, or ``None`` to let the caller fall back to
    the plain ``stdout`` text pass (byte-identical text, confidence unknown). Never
    raises; any failure yields ``None`` so text extraction is never lost.
    """
    tmpdir = None
    try:
        tmpdir = tempfile.mkdtemp(prefix="khy-ocr-cli-")
        outbase = os.path.join(tmpdir, "out")
        proc = subprocess.run(
            [bin_path, image_path, outbase, "-l", eff_lang, "txt", "tsv"],
            capture_output=True, text=True, timeout=30,
        )
        txt_path = outbase + ".txt"
        if proc.returncode != 0 or not os.path.isfile(txt_path):
            return None
        with open(txt_path, "r", encoding="utf-8", errors="replace") as fh:
            text = fh.read().strip()
        if not text:
            err = (proc.stderr or "tesseract produced no text").strip()
            return {"success": False, "error": err, "needsAiFallback": True}
        result = {
            "success": True,
            "text": text,
            "lang": eff_lang,
            "engine": "tesseract-cli",
        }
        avg_conf = _mean_tsv_confidence(outbase + ".tsv")
        if avg_conf is None:
            # No usable confidence signal → do not fabricate one (byte-revert to
            # the historical CLI contract: confidence unknown, no low-conf flag).
            result["confidence"] = 0
            result["needsAiFallback"] = False
        else:
            result["confidence"] = round(avg_conf, 1)
            result["needsAiFallback"] = avg_conf < 60
        return result
    except Exception:  # noqa: BLE001 — fall back to the plain stdout pass
        return None
    finally:
        if tmpdir:
            shutil.rmtree(tmpdir, ignore_errors=True)


def _orient_gate_enabled():
    """KHY_OCR_AUTO_ORIENT — default-on. Off-words {0,false,off,no} byte-revert to
    the historical no-reorientation behavior. Read from the environment so a
    text-only deployment can disable the extra rotation passes if desired."""
    val = os.environ.get("KHY_OCR_AUTO_ORIENT")
    if val is None:
        return True
    return str(val).strip().lower() not in ("0", "false", "off", "no")


def _maybe_reorient(bin_path, image_path, eff_lang, base):
    """Recover text from a rotated image when the upright pass is weak.

    tesseract's own OSD (``--psm 0``) is unreliable on sparse text ("Too few
    characters. Skipping this page"), so instead of trusting OSD we brute-force
    the three non-zero orientations (90/180/270), OCR each with the SAME confident
    txt+tsv pass, and keep the best-scoring readable result. This is the first
    *corrective* OCR axis: the confidence axis merely warns that a sideways photo
    produced garbage (a rotated page still yields a deceptively high ~51 conf); this
    actually restores the upright text so a non-vision model can read it.

    Guardrails (never degrade the honest baseline):
      * only runs when the gate is on AND ``base`` is weak (failed or
        needsAiFallback) — an already-confident upright read is left untouched;
      * requires PIL to rotate; absent PIL → return ``base`` unchanged;
      * a rotated candidate is accepted only when it SUCCEEDS, clears the
        confidence floor (>=60), and beats the base confidence by a real margin —
        otherwise the original ``base`` is returned verbatim;
      * fail-soft: any exception → ``base`` (text extraction is never lost).

    On acceptance the winning dict carries ``orientationCorrected`` = the degrees
    the image was rotated (90/180/270); the untouched base keeps it at 0.
    """
    if not _orient_gate_enabled():
        return base
    # Attempt recovery unless the upright read is already CLEARLY good. A rotated
    # page yields deceptively-moderate confidence (a 2-line invoice reads as
    # ~62 garbage — above the <60 low-confidence flag), so gating on
    # needsAiFallback alone misses it. Only a genuinely strong upright read
    # (>= _ORIENT_SKIP_CONF and not flagged) skips the extra passes.
    _ORIENT_SKIP_CONF = 80.0
    base_ok = isinstance(base, dict) and base.get("success") is True
    base_conf = 0.0
    if base_ok:
        try:
            base_conf = float(base.get("confidence") or 0)
        except (TypeError, ValueError):
            base_conf = 0.0
        if base.get("needsAiFallback") is not True and base_conf >= _ORIENT_SKIP_CONF:
            return base
    tmpdir = None
    try:
        from PIL import Image  # optional dependency; absent → no reorientation
    except Exception:  # noqa: BLE001 — PIL missing → keep base
        return base
    try:
        tmpdir = tempfile.mkdtemp(prefix="khy-ocr-orient-")
        best = None
        best_deg = 0
        for deg in (90, 180, 270):
            try:
                # PIL rotate() is counter-clockwise; expand keeps the whole frame.
                rotated_path = os.path.join(tmpdir, "r%d.png" % deg)
                with Image.open(image_path) as im:
                    im.rotate(-deg, expand=True).save(rotated_path)
            except Exception:  # noqa: BLE001 — this angle failed → skip it
                continue
            cand = _ocr_via_cli_with_confidence(bin_path, rotated_path, eff_lang)
            if not (isinstance(cand, dict) and cand.get("success") is True):
                continue
            try:
                cand_conf = float(cand.get("confidence") or 0)
            except (TypeError, ValueError):
                continue
            if best is None or cand_conf > float(best.get("confidence") or 0):
                best = cand
                best_deg = deg
        if best is None:
            return base
        best_conf = float(best.get("confidence") or 0)
        # Accept only a DECISIVELY-better, confident rotation: it must clear the
        # confidence floor (>=60) AND beat the upright read by a wide margin. A
        # wide margin (20) makes a false correction — a rotation coincidentally
        # edging out a correct-but-low-quality upright read — vanishingly unlikely,
        # while a genuinely misoriented page beats its garbage baseline by ~30+.
        if best_conf < 60 or best_conf < base_conf + 20:
            return base
        best["orientationCorrected"] = best_deg
        best["lang"] = eff_lang
        return best
    except Exception:  # noqa: BLE001 — any failure → keep the honest base
        return base
    finally:
        if tmpdir:
            shutil.rmtree(tmpdir, ignore_errors=True)


def _upscale_gate_enabled():
    """KHY_OCR_UPSCALE — default-on. Off-words {0,false,off,no} byte-revert to the
    historical single-scale behavior. Lets a deployment disable the extra enlarge
    passes if the CPU cost of re-OCRing an upscaled image is unwanted."""
    val = os.environ.get("KHY_OCR_UPSCALE")
    if val is None:
        return True
    return str(val).strip().lower() not in ("0", "false", "off", "no")


def _maybe_upscale(bin_path, image_path, eff_lang, base):
    """Recover text from a LOW-RESOLUTION image by enlarging it before OCR.

    tesseract wants ~300 DPI; glyphs only a handful of pixels tall read as empty
    or garbage at native size, yet a plain LANCZOS upscale reconstructs enough
    edge information to recognize them (measured: a 102x10 'INVOICE' crop returns
    nothing natively but reads at conf ~96 when enlarged 2x — while 3x happened to
    miss, so a single fixed factor is unreliable). This is the second *corrective*
    OCR axis (sibling to ``_maybe_reorient``): the confidence/coverage axes merely
    note that a tiny image produced little/no text; this actually enlarges it so a
    non-vision model can read the recovered text.

    Guardrails (never degrade the honest baseline, never blow up memory):
      * gate KHY_OCR_UPSCALE (default-on); off → return ``base`` unchanged;
      * only runs when ``base`` is WEAK (failed / needsAiFallback / low conf) — a
        clearly-good read (>= _UPSCALE_SKIP_CONF and not flagged) is left untouched;
      * only enlarges genuinely SMALL sources (max side < _UPSCALE_MAX_SRC) and
        skips any factor whose result would exceed _UPSCALE_MAX_DST;
      * requires PIL to resize; absent PIL → return ``base``;
      * brute-forces 2x/3x/4x (one fixed factor is unreliable) and keeps the
        best-scoring readable result; a candidate is accepted only when it clears
        the confidence floor (>=60), and — when the base already had text — also
        beats it by a wide margin, so a coincidental tie never rewrites a
        correct-but-modest native read; when the base was EMPTY any confident
        recovery is a pure win;
      * fail-soft: any exception → ``base`` (text extraction is never lost).

    On acceptance the winning dict carries ``upscaledFactor`` = the integer scale
    applied (2/3/4); an untouched base keeps it at 0.
    """
    if not _upscale_gate_enabled():
        return base
    _UPSCALE_SKIP_CONF = 80.0
    _UPSCALE_MAX_SRC = 1000  # only enlarge small/low-res images
    _UPSCALE_MAX_DST = 4000  # never produce an enormous canvas
    base_ok = isinstance(base, dict) and base.get("success") is True
    base_conf = 0.0
    base_has_text = False
    if base_ok:
        try:
            base_conf = float(base.get("confidence") or 0)
        except (TypeError, ValueError):
            base_conf = 0.0
        base_has_text = bool(base.get("text"))
        # Clearly-good read → don't perturb.
        if base.get("needsAiFallback") is not True and base_conf >= _UPSCALE_SKIP_CONF:
            return base
    tmpdir = None
    try:
        from PIL import Image  # optional dependency; absent → no upscaling
    except Exception:  # noqa: BLE001 — PIL missing → keep base
        return base
    try:
        with Image.open(image_path) as im:
            src_w, src_h = im.size
        if max(src_w, src_h) >= _UPSCALE_MAX_SRC:
            return base  # already high-res; enlarging won't help, only costs memory
        tmpdir = tempfile.mkdtemp(prefix="khy-ocr-upscale-")
        best = None
        best_factor = 0
        for factor in (2, 3, 4):
            if max(src_w, src_h) * factor > _UPSCALE_MAX_DST:
                continue
            try:
                up_path = os.path.join(tmpdir, "u%d.png" % factor)
                with Image.open(image_path) as im:
                    im.resize((src_w * factor, src_h * factor), Image.LANCZOS).save(up_path)
            except Exception:  # noqa: BLE001 — this factor failed → skip it
                continue
            cand = _ocr_via_cli_with_confidence(bin_path, up_path, eff_lang)
            if not (isinstance(cand, dict) and cand.get("success") is True):
                continue
            try:
                cand_conf = float(cand.get("confidence") or 0)
            except (TypeError, ValueError):
                continue
            if best is None or cand_conf > float(best.get("confidence") or 0):
                best = cand
                best_factor = factor
        if best is None:
            return base
        best_conf = float(best.get("confidence") or 0)
        if best_conf < 60:
            return base
        if base_ok and base_has_text and best_conf < base_conf + 20:
            return base
        best["upscaledFactor"] = best_factor
        best["lang"] = eff_lang
        return best
    except Exception:  # noqa: BLE001 — any failure → keep the honest base
        return base
    finally:
        if tmpdir:
            shutil.rmtree(tmpdir, ignore_errors=True)


def _ocr_via_cli(image_path, lang):
    """OCR an image by invoking the tesseract CLI binary directly.

    pytesseract is only a thin wrapper around this same binary; when the wrapper
    (or PIL) is not installed we can still extract text as long as the tesseract
    engine itself is on PATH. This keeps khy's image-recognition fallback working
    for text-only models without requiring the optional Python doc dependencies.

    Prefers a single txt+tsv pass (honest confidence + needsAiFallback); if that
    is unavailable it degrades to a plain ``stdout`` text pass (byte-identical
    text, confidence unknown). Returns None when the tesseract binary is
    unavailable (so the caller can emit a precise "install tesseract" message),
    otherwise a result dict.
    """
    bin_path = shutil.which("tesseract")
    if not bin_path:
        return None
    eff_lang = _resolve_lang(lang, _tesseract_available_langs(bin_path))
    honest = _ocr_via_cli_with_confidence(bin_path, image_path, eff_lang)
    if honest is not None:
        # requestedLang = the caller's original 'a+b+c' spec; lang = the subset
        # actually usable on this box. When they differ, a language pack was
        # missing and text in the dropped languages was NOT recognized — the JS
        # side surfaces an honest caveat so a text-only model isn't fed silently
        # un-recognized script as if it were authoritative.
        honest["requestedLang"] = lang
        # Corrective axis: if the upright read is weak (a sideways photo yields
        # deceptively-confident garbage), recover the best rotation. Untouched
        # base keeps orientationCorrected=0. Fail-soft — never loses text.
        honest = _maybe_reorient(bin_path, image_path, eff_lang, honest)
        # Second corrective axis: if the read is still weak because the image is
        # tiny/low-res, enlarge it (2x/3x/4x, keep best) and re-OCR. Untouched base
        # keeps upscaledFactor=0. Fail-soft — never loses text.
        honest = _maybe_upscale(bin_path, image_path, eff_lang, honest)
        honest["requestedLang"] = lang
        if "orientationCorrected" not in honest:
            honest["orientationCorrected"] = 0
        if "upscaledFactor" not in honest:
            honest["upscaledFactor"] = 0
        return honest
    # Degraded path: plain stdout text pass (byte-identical text, no confidence).
    try:
        proc = subprocess.run(
            [bin_path, image_path, "stdout", "-l", eff_lang],
            capture_output=True, text=True, timeout=30,
        )
    except Exception as e:  # noqa: BLE001 — report any spawn/timeout failure
        return {
            "success": False,
            "error": f"tesseract CLI failed: {str(e)}",
            "needsAiFallback": True,
        }
    text = (proc.stdout or "").strip()
    if not text:
        err = (proc.stderr or "tesseract produced no text").strip()
        return {"success": False, "error": err, "needsAiFallback": True}
    return {
        "success": True,
        "text": text,
        "confidence": 0,  # CLI text pass does not compute per-word confidence
        "lang": eff_lang,
        "requestedLang": lang,
        "engine": "tesseract-cli",
        "needsAiFallback": False,
        "orientationCorrected": 0,  # degraded pass does not attempt reorientation
        "upscaledFactor": 0,  # degraded pass does not attempt upscaling
    }


def ocr_image(image_path, lang="chi_sim+eng"):
    """Extract text from image using Tesseract OCR.

    Preferred path uses the pytesseract wrapper (which also yields a confidence
    score). When the wrapper or PIL is missing, we fall back to driving the
    tesseract CLI binary directly so OCR — and therefore image recognition for
    text-only models — keeps working with just the engine installed.
    """
    if not os.path.isfile(image_path):
        _output({"success": False, "error": f"Image not found: {image_path}"})
        return

    size = os.path.getsize(image_path)
    if size > MAX_IMAGE_SIZE:
        _output({
            "success": False,
            "error": f"Image too large: {size / (1024 * 1024):.1f}MB (max 20MB)",
        })
        return

    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        cli = _ocr_via_cli(image_path, lang)
        if cli is None:
            _output({
                "success": False,
                "error": "Tesseract OCR engine not found. Install: brew install tesseract (macOS) / apt install tesseract-ocr (Linux) / choco install tesseract (Windows). For higher accuracy also: pip install khy-os[doc]",
                "needsAiFallback": True,
            })
        else:
            _output(cli)
        return

    eff_lang = _resolve_lang(lang, _tesseract_available_langs(shutil.which("tesseract") or "tesseract"))
    try:
        img = Image.open(image_path)
        text = pytesseract.image_to_string(img, lang=eff_lang)

        # Calculate average confidence
        conf_data = pytesseract.image_to_data(
            img, lang=eff_lang, output_type=pytesseract.Output.DICT
        )
        confs = [
            int(c)
            for c in conf_data.get("conf", [])
            if str(c).lstrip("-").isdigit() and int(c) > 0
        ]
        avg_conf = sum(confs) / len(confs) if confs else 0

        base = {
            "success": True,
            "text": text.strip(),
            "confidence": round(avg_conf, 1),
            "lang": eff_lang,
            "requestedLang": lang,
            "needsAiFallback": avg_conf < 60,
        }
        # Corrective axis: a weak upright read (sideways photo → confident garbage)
        # is retried at 90/180/270 via the CLI txt+tsv pass; the best confident
        # rotation wins. Untouched base keeps orientationCorrected=0. Fail-soft.
        base = _maybe_reorient(shutil.which("tesseract") or "tesseract", image_path, eff_lang, base)
        # Second corrective axis: a tiny/low-res image is enlarged (2x/3x/4x, keep
        # best) and re-OCRed. Untouched base keeps upscaledFactor=0. Fail-soft.
        base = _maybe_upscale(shutil.which("tesseract") or "tesseract", image_path, eff_lang, base)
        if "orientationCorrected" not in base:
            base["orientationCorrected"] = 0
        if "upscaledFactor" not in base:
            base["upscaledFactor"] = 0
        _output(base)
    except pytesseract.TesseractNotFoundError:
        # Wrapper is present but cannot locate the binary — try the CLI path
        # (covers PATH quirks) before giving up.
        cli = _ocr_via_cli(image_path, lang)
        if cli is not None:
            _output(cli)
            return
        _output({
            "success": False,
            "error": "Tesseract OCR engine not found. Install: brew install tesseract (macOS) / apt install tesseract-ocr (Linux) / choco install tesseract (Windows)",
            "needsAiFallback": True,
        })
    except Exception as e:
        _output({
            "success": False,
            "error": f"OCR failed: {str(e)}",
            "needsAiFallback": True,
        })


def text_to_docx(text, output_path):
    """Write recognized text into an editable Word (.docx) document."""
    try:
        from docx import Document
    except ImportError:
        _output({
            "success": False,
            "error": "python-docx not available. It is normally installed with pdf2docx. Run: pip install khy-os[doc]",
        })
        return

    if not text or not text.strip():
        _output({"success": False, "error": "No text to write"})
        return

    try:
        out_dir = os.path.dirname(output_path)
        if out_dir and not os.path.isdir(out_dir):
            os.makedirs(out_dir, exist_ok=True)

        doc = Document()
        for paragraph in text.split("\n"):
            doc.add_paragraph(paragraph)
        doc.save(output_path)

        out_size = os.path.getsize(output_path)
        _output({
            "success": True,
            "output": output_path,
            "outputSize": out_size,
            "message": f"Saved: {os.path.basename(output_path)} ({out_size / 1024:.0f}KB)",
        })
    except Exception as e:
        _output({"success": False, "error": f"Failed to create Word document: {str(e)}"})


# Built-in Word style names that denote a document title / heading / caption,
# including the localized Chinese names Word uses for zh-CN documents. Title
# styling commonly needs both because the user's docs are often Chinese.
TITLE_STYLE_NAMES = {
    "Title", "Heading 1", "Heading 2", "Caption",
    "标题", "标题 1", "标题 2", "题注",
}


def _normalize_hex(color):
    """Return a 6-char uppercase hex string (no leading '#') or None."""
    if not color:
        return None
    s = str(color).strip().lstrip("#").upper()
    if len(s) == 3:  # shorthand like F00 → FF0000
        s = "".join(ch * 2 for ch in s)
    if len(s) != 6 or any(c not in "0123456789ABCDEF" for c in s):
        return None
    return s


def title_style(input_path, output_path, match=None, style=None,
                size_pt=None, color=None):
    """Restyle a Word document's title/heading/caption: set font size and/or
    color on the matching paragraphs.

    Targeting (any paragraph matching is restyled):
      - style: explicit style name (e.g. 'Title' / 'Heading 1' / '标题'); if
        omitted, the built-in TITLE_STYLE_NAMES set is used.
      - match: exact paragraph text (stripped) to target a specific heading.
    At least one of size_pt / color must be provided.

    Font size and color MUST be set on each run (not the paragraph), and a
    heading's text may be split across several runs, so every run is updated.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        _output({
            "success": False,
            "needsDep": True,
            "error": "python-docx not installed. Run: pip install khy-os[doc]",
            "hint": "pip install khy-os[doc]",
        })
        return

    if not os.path.isfile(input_path):
        _output({"success": False, "error": f"File not found: {input_path}"})
        return

    # Validate styling intent up front so we fail clean, not mid-document.
    pt_value = None
    if size_pt is not None and str(size_pt) != "":
        try:
            pt_value = float(size_pt)
            if pt_value <= 0 or pt_value > 1638:  # Word's practical Pt ceiling
                raise ValueError
        except (TypeError, ValueError):
            _output({"success": False, "error": f"Invalid size (pt): {size_pt}"})
            return

    hex_value = None
    if color is not None and str(color) != "":
        hex_value = _normalize_hex(color)
        if hex_value is None:
            _output({"success": False,
                     "error": f"Invalid color (need 3/6 hex digits): {color}"})
            return

    if pt_value is None and hex_value is None:
        _output({"success": False,
                 "error": "Nothing to change: provide size (pt) and/or color (hex)."})
        return

    target_styles = {style} if style else TITLE_STYLE_NAMES
    target_text = match.strip() if match else None

    try:
        doc = Document(input_path)

        def _matches(paragraph):
            if target_text is not None:
                return paragraph.text.strip() == target_text
            name = paragraph.style.name if paragraph.style else None
            return name in target_styles

        changed_runs = 0
        matched_paragraphs = 0
        for paragraph in doc.paragraphs:
            if not _matches(paragraph):
                continue
            matched_paragraphs += 1
            for run in paragraph.runs:
                if pt_value is not None:
                    run.font.size = Pt(pt_value)
                if hex_value is not None:
                    run.font.color.rgb = RGBColor.from_string(hex_value)
                changed_runs += 1

        if matched_paragraphs == 0:
            _output({
                "success": False,
                "changed": 0,
                "matchedParagraphs": 0,
                "error": "No matching title/heading found.",
                "hint": "Try a different --match text or --style "
                        "(e.g. 'Title' / 'Heading 1' / '标题').",
            })
            return

        out_dir = os.path.dirname(output_path)
        if out_dir and not os.path.isdir(out_dir):
            os.makedirs(out_dir, exist_ok=True)
        doc.save(output_path)

        out_size = os.path.getsize(output_path)
        _output({
            "success": True,
            "output": output_path,
            "matchedParagraphs": matched_paragraphs,
            "changed": changed_runs,
            "sizePt": pt_value,
            "color": hex_value,
            "outputSize": out_size,
            "message": (
                f"Restyled {matched_paragraphs} title paragraph(s), "
                f"{changed_runs} run(s) → {os.path.basename(output_path)}"
            ),
        })
    except Exception as e:
        _output({"success": False,
                 "error": f"Failed to restyle title: {str(e)}"})


# Image extensions Pillow can open and merge into a PDF.
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tiff", ".tif", ".webp"}


def images_to_pdf(output_path, image_paths):
    """Merge one or more images into a single (multi-page) PDF via Pillow.

    Each image is loaded and converted to RGB first — PNGs with an alpha
    channel or a palette cannot be saved to PDF directly, which is the most
    common "image to PDF failed" pitfall. The first image carries the rest as
    appended pages, so N images → one N-page PDF.
    """
    try:
        from PIL import Image
    except ImportError:
        _output({
            "success": False,
            "needsDep": True,
            "error": "Pillow not installed. Run: pip install khy-os[doc]",
            "hint": "pip install khy-os[doc]",
        })
        return

    if not image_paths:
        _output({"success": False, "error": "No input images provided"})
        return

    frames = []
    try:
        for p in image_paths:
            if not os.path.isfile(p):
                _output({"success": False, "error": f"Image not found: {p}"})
                return
            size = os.path.getsize(p)
            if size > MAX_IMAGE_SIZE:
                _output({
                    "success": False,
                    "error": f"Image too large: {os.path.basename(p)} "
                             f"{size / (1024 * 1024):.1f}MB (max 20MB)",
                })
                return
            img = Image.open(p)
            # PDF has no alpha/palette concept — flatten to RGB.
            frames.append(img.convert("RGB"))

        out_dir = os.path.dirname(output_path)
        if out_dir and not os.path.isdir(out_dir):
            os.makedirs(out_dir, exist_ok=True)

        first, rest = frames[0], frames[1:]
        first.save(output_path, "PDF", save_all=True, append_images=rest)

        out_size = os.path.getsize(output_path)
        _output({
            "success": True,
            "output": output_path,
            "pages": len(frames),
            "inputs": len(image_paths),
            "outputSize": out_size,
            "message": (
                f"Merged {len(frames)} image(s) → "
                f"{os.path.basename(output_path)} ({out_size / 1024:.0f}KB)"
            ),
        })
    except Exception as e:
        _output({"success": False, "error": f"Image-to-PDF failed: {str(e)}"})


def pdf_to_text(input_path, output_path):
    """Extract the text layer of a PDF to a .txt file via pypdf.

    A PDF that is really a scan has no text layer; extraction yields an empty
    string. Rather than write an empty file and claim success, we fail with a
    hint pointing the user at OCR (image-to-txt).
    """
    try:
        from pypdf import PdfReader
    except ImportError:
        _output({
            "success": False,
            "needsDep": True,
            "error": "pypdf not installed. Run: pip install khy-os[doc]",
            "hint": "pip install khy-os[doc]",
        })
        return

    if not os.path.isfile(input_path):
        _output({"success": False, "error": f"File not found: {input_path}"})
        return

    size = os.path.getsize(input_path)
    if size > MAX_PDF_SIZE:
        _output({
            "success": False,
            "error": f"PDF too large: {size / (1024 * 1024):.1f}MB (max 50MB)",
        })
        return

    try:
        reader = PdfReader(input_path)
        pages = [(page.extract_text() or "") for page in reader.pages]
        text = "\n\n".join(pages).strip()

        if not text:
            _output({
                "success": False,
                "error": "No extractable text layer found in this PDF.",
                "hint": "The PDF is likely a scan. Convert the page images to "
                        "text with OCR instead (image → txt).",
                "needsAiFallback": True,
            })
            return

        out_dir = os.path.dirname(output_path)
        if out_dir and not os.path.isdir(out_dir):
            os.makedirs(out_dir, exist_ok=True)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(text)

        _output({
            "success": True,
            "output": output_path,
            "pages": len(pages),
            "chars": len(text),
            "message": (
                f"Extracted {len(text)} chars from {len(pages)} page(s) → "
                f"{os.path.basename(output_path)}"
            ),
        })
    except Exception as e:
        _output({"success": False, "error": f"PDF-to-text failed: {str(e)}"})


def docx_to_text(input_path, output_path):
    """Extract the paragraph text of a Word (.docx) document to a .txt file."""
    try:
        from docx import Document
    except ImportError:
        _output({
            "success": False,
            "needsDep": True,
            "error": "python-docx not installed. Run: pip install khy-os[doc]",
            "hint": "pip install khy-os[doc]",
        })
        return

    if not os.path.isfile(input_path):
        _output({"success": False, "error": f"File not found: {input_path}"})
        return

    try:
        doc = Document(input_path)
        text = "\n".join(p.text for p in doc.paragraphs).strip()

        out_dir = os.path.dirname(output_path)
        if out_dir and not os.path.isdir(out_dir):
            os.makedirs(out_dir, exist_ok=True)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(text)

        _output({
            "success": True,
            "output": output_path,
            "chars": len(text),
            "message": (
                f"Extracted {len(text)} chars → "
                f"{os.path.basename(output_path)}"
            ),
        })
    except Exception as e:
        _output({"success": False, "error": f"Word-to-text failed: {str(e)}"})


def text_file_to_docx(input_path, output_path):
    """Convert a plain-text file into an editable Word (.docx) document.

    Reads the text from a FILE (one paragraph per line) rather than taking it
    on argv, so large documents are not pushed through the command line.
    """
    try:
        from docx import Document
    except ImportError:
        _output({
            "success": False,
            "needsDep": True,
            "error": "python-docx not installed. Run: pip install khy-os[doc]",
            "hint": "pip install khy-os[doc]",
        })
        return

    if not os.path.isfile(input_path):
        _output({"success": False, "error": f"File not found: {input_path}"})
        return

    try:
        with open(input_path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()

        out_dir = os.path.dirname(output_path)
        if out_dir and not os.path.isdir(out_dir):
            os.makedirs(out_dir, exist_ok=True)

        doc = Document()
        for line in text.split("\n"):
            doc.add_paragraph(line)
        doc.save(output_path)

        out_size = os.path.getsize(output_path)
        _output({
            "success": True,
            "output": output_path,
            "outputSize": out_size,
            "message": (
                f"Saved: {os.path.basename(output_path)} "
                f"({out_size / 1024:.0f}KB)"
            ),
        })
    except Exception as e:
        _output({"success": False, "error": f"Text-to-Word failed: {str(e)}"})


if __name__ == "__main__":
    if len(sys.argv) < 2:
        _output({"success": False, "error": "Usage: docHelper.py <check|pdf2word|ocr|text2docx|title-style|img2pdf|pdf2txt|docx2txt|txt2docx> [args]"})
        sys.exit(1)

    cmd = sys.argv[1]

    if cmd == "check":
        check_deps()
    elif cmd == "pdf2word":
        if len(sys.argv) < 4:
            _output({"success": False, "error": "Usage: docHelper.py pdf2word <input.pdf> <output.docx>"})
            sys.exit(1)
        pdf_to_word(sys.argv[2], sys.argv[3])
    elif cmd == "ocr":
        if len(sys.argv) < 3:
            _output({"success": False, "error": "Usage: docHelper.py ocr <image_path> [lang]"})
            sys.exit(1)
        lang = sys.argv[3] if len(sys.argv) > 3 else "chi_sim+eng"
        ocr_image(sys.argv[2], lang)
    elif cmd == "text2docx":
        if len(sys.argv) < 4:
            _output({"success": False, "error": "Usage: docHelper.py text2docx <text> <output.docx>"})
            sys.exit(1)
        text_to_docx(sys.argv[2], sys.argv[3])
    elif cmd == "title-style":
        # Usage: docHelper.py title-style <input.docx> <output.docx>
        #          [--match <text>] [--style <name>] [--size <pt>] [--color <hex>]
        if len(sys.argv) < 4:
            _output({"success": False, "error": "Usage: docHelper.py title-style <input.docx> <output.docx> [--match T] [--style S] [--size PT] [--color HEX]"})
            sys.exit(1)
        opts = {"match": None, "style": None, "size": None, "color": None}
        rest = sys.argv[4:]
        i = 0
        while i < len(rest):
            flag = rest[i]
            key = flag[2:] if flag.startswith("--") else None
            if key in opts and i + 1 < len(rest):
                opts[key] = rest[i + 1]
                i += 2
            else:
                i += 1
        title_style(sys.argv[2], sys.argv[3],
                    match=opts["match"], style=opts["style"],
                    size_pt=opts["size"], color=opts["color"])
    elif cmd == "img2pdf":
        # Usage: docHelper.py img2pdf <output.pdf> <img1> [img2 ...]
        if len(sys.argv) < 4:
            _output({"success": False, "error": "Usage: docHelper.py img2pdf <output.pdf> <img1> [img2 ...]"})
            sys.exit(1)
        images_to_pdf(sys.argv[2], sys.argv[3:])
    elif cmd == "pdf2txt":
        if len(sys.argv) < 4:
            _output({"success": False, "error": "Usage: docHelper.py pdf2txt <input.pdf> <output.txt>"})
            sys.exit(1)
        pdf_to_text(sys.argv[2], sys.argv[3])
    elif cmd == "docx2txt":
        if len(sys.argv) < 4:
            _output({"success": False, "error": "Usage: docHelper.py docx2txt <input.docx> <output.txt>"})
            sys.exit(1)
        docx_to_text(sys.argv[2], sys.argv[3])
    elif cmd == "txt2docx":
        if len(sys.argv) < 4:
            _output({"success": False, "error": "Usage: docHelper.py txt2docx <input.txt> <output.docx>"})
            sys.exit(1)
        text_file_to_docx(sys.argv[2], sys.argv[3])
    else:
        _output({"success": False, "error": f"Unknown command: {cmd}"})
        sys.exit(1)
