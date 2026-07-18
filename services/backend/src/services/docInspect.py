#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
docInspect.py — read-only document **format** inspector.

The counterpart to docTypeset.py (which WRITES docx). This module READS an existing
document and reports its precise formatting so khy can answer questions like
"what font / size / first-line indent / line spacing / heading-1 style does this use?".

Contract (mirrors docTypeset.py / docHelper.py):
    argv:  docInspect.py inspect <path>
    stdout: a single JSON object. On any failure: {"success": false, "error": ...}
            (+ "needsDep": true when a Python lib is missing — with an install hint).
    Never crashes the caller; all imports are lazy and every path is guarded.

Supported:
    .docx  → python-docx: page size/margins, per-paragraph style/font/size/bold/
             alignment/first-line-indent(chars)/line-spacing, heading outline,
             and a dominant-body-font summary.
    .pdf   → PyMuPDF (fitz) if available: per-page dominant font/size + spans sample.
             If fitz is missing, fail-soft with an install hint (never crash).
    .md/.markdown/.txt/source → structural outline (headings, line/paragraph counts).
             Visual formatting (font/size) is N/A for plain text and reported as such.
"""

import sys
import json
import os

EMU_PER_CM = 360000.0
EMU_PER_PT = 12700.0


def _output(data):
    print(json.dumps(data, ensure_ascii=False))


def _err(msg, **extra):
    out = {"success": False, "error": msg}
    out.update(extra)
    _output(out)


def _emu_to_cm(emu):
    try:
        return round(float(emu) / EMU_PER_CM, 3)
    except (TypeError, ValueError):
        return None


def _emu_to_pt(emu):
    try:
        return round(float(emu) / EMU_PER_PT, 2)
    except (TypeError, ValueError):
        return None


# ── docx ───────────────────────────────────────────────────────────────────

def _run_east_asia(run, qn):
    """Read the eastAsia font from a run's XML (python-docx does not expose it)."""
    try:
        rpr = run._element.find(qn("w:rPr"))
        if rpr is None:
            return None
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            return None
        return rfonts.get(qn("w:eastAsia"))
    except Exception:
        return None


def _style_east_asia(style, qn):
    """Read the eastAsia font declared on a style's own rPr (style-level default)."""
    try:
        el = style.element
        rpr = el.find(qn("w:rPr"))
        if rpr is None:
            return None
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            return None
        return rfonts.get(qn("w:eastAsia"))
    except Exception:
        return None


def _style_font(style, qn):
    """
    Effective style-level font/size, walking the style inheritance chain
    (e.g. Heading 1 → Heading base → Normal) so headings whose formatting lives
    in the style definition (not the run) are still reported.
    """
    ascii_name = None
    east = None
    size_pt = None
    seen = 0
    s = style
    while s is not None and seen < 12:
        seen += 1
        try:
            if ascii_name is None and s.font is not None and s.font.name:
                ascii_name = s.font.name
            if size_pt is None and s.font is not None and s.font.size is not None:
                size_pt = _emu_to_pt(int(s.font.size))
        except Exception:
            pass
        if east is None:
            east = _style_east_asia(s, qn)
        if ascii_name and east and size_pt:
            break
        try:
            s = s.base_style
        except Exception:
            break
    return ascii_name, east, size_pt


def _inspect_docx(path):
    from docx import Document
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: F401

    doc = Document(path)

    # Page geometry
    page = {}
    try:
        sec = doc.sections[0]
        page = {
            "widthCm": _emu_to_cm(sec.page_width),
            "heightCm": _emu_to_cm(sec.page_height),
            "marginTopCm": _emu_to_cm(sec.top_margin),
            "marginBottomCm": _emu_to_cm(sec.bottom_margin),
            "marginLeftCm": _emu_to_cm(sec.left_margin),
            "marginRightCm": _emu_to_cm(sec.right_margin),
        }
        w, h = page.get("widthCm"), page.get("heightCm")
        # A4 = 21.0 x 29.7 cm
        page["isA4"] = bool(w and h and abs(w - 21.0) < 0.2 and abs(h - 29.7) < 0.2)
    except Exception as e:
        page = {"error": "page read failed: %s" % e}

    paragraphs = []
    outline = []
    body_font_counter = {}
    body_size_counter = {}
    MAX_PARA = 800

    for i, para in enumerate(doc.paragraphs):
        if i >= MAX_PARA:
            break
        text = para.text or ""
        style_name = ""
        try:
            style_name = para.style.name if para.style is not None else ""
        except Exception:
            style_name = ""

        # Paragraph format
        pf = para.paragraph_format
        first_line_emu = None
        line_spacing = None
        line_spacing_rule = None
        align = None
        try:
            first_line_emu = int(pf.first_line_indent) if pf.first_line_indent is not None else None
        except Exception:
            first_line_emu = None
        try:
            ls = pf.line_spacing
            if ls is not None:
                # float → multiple; Length → exact (pt)
                line_spacing = float(ls) if isinstance(ls, float) else _emu_to_pt(int(ls))
        except Exception:
            line_spacing = None
        try:
            line_spacing_rule = str(pf.line_spacing_rule) if pf.line_spacing_rule is not None else None
        except Exception:
            line_spacing_rule = None
        try:
            align = str(para.alignment) if para.alignment is not None else None
        except Exception:
            align = None

        # First run font (representative of the paragraph body)
        font_ascii = None
        font_east = None
        size_pt = None
        bold = None
        if para.runs:
            r0 = para.runs[0]
            try:
                font_ascii = r0.font.name
            except Exception:
                font_ascii = None
            font_east = _run_east_asia(r0, qn)
            try:
                size_pt = _emu_to_pt(int(r0.font.size)) if r0.font.size is not None else None
            except Exception:
                size_pt = None
            try:
                bold = bool(r0.font.bold) if r0.font.bold is not None else None
            except Exception:
                bold = None

        # Fall back to the style definition when run-level formatting is absent
        # (headings/Title usually carry their font+size on the style, not the run).
        if (font_ascii is None or font_east is None or size_pt is None) and para.style is not None:
            s_ascii, s_east, s_size = _style_font(para.style, qn)
            if font_ascii is None:
                font_ascii = s_ascii
            if font_east is None:
                font_east = s_east
            if size_pt is None:
                size_pt = s_size

        # First-line indent expressed in CJK characters (chars ≈ indent_pt / size_pt)
        first_line_chars = None
        if first_line_emu and size_pt:
            try:
                first_line_chars = round((_emu_to_pt(first_line_emu) or 0) / float(size_pt), 1)
            except Exception:
                first_line_chars = None

        is_heading = style_name.lower().startswith("heading") or style_name in ("Title",)
        if is_heading and text.strip():
            level = 0
            digits = "".join(ch for ch in style_name if ch.isdigit())
            if digits:
                try:
                    level = int(digits)
                except ValueError:
                    level = 0
            outline.append({"level": level or (1 if style_name == "Title" else 0),
                            "style": style_name, "text": text.strip()[:120],
                            "fontEastAsia": font_east, "fontAscii": font_ascii, "sizePt": size_pt, "bold": bold})
        elif text.strip():
            # Tally dominant body font/size
            key_f = font_east or font_ascii
            if key_f:
                body_font_counter[key_f] = body_font_counter.get(key_f, 0) + 1
            if size_pt:
                body_size_counter[size_pt] = body_size_counter.get(size_pt, 0) + 1

        paragraphs.append({
            "index": i + 1,
            "style": style_name,
            "isHeading": is_heading,
            "text": text[:200],
            "fontAscii": font_ascii,
            "fontEastAsia": font_east,
            "sizePt": size_pt,
            "bold": bold,
            "align": align,
            "firstLineIndentCm": _emu_to_cm(first_line_emu) if first_line_emu else None,
            "firstLineIndentChars": first_line_chars,
            "lineSpacing": line_spacing,
            "lineSpacingRule": line_spacing_rule,
        })

    def _dominant(counter):
        if not counter:
            return None
        return max(counter.items(), key=lambda kv: kv[1])[0]

    # Heading-1 summary (first H1)
    h1 = next((o for o in outline if o["level"] == 1), None)

    summary = {
        "bodyFont": _dominant(body_font_counter),
        "bodySizePt": _dominant(body_size_counter),
        "heading1": ({"font": h1["fontEastAsia"] or h1["fontAscii"], "sizePt": h1["sizePt"], "bold": h1["bold"]}
                     if h1 else None),
        "paragraphCount": len(doc.paragraphs),
        "headingCount": len(outline),
    }

    return {
        "success": True,
        "format": "docx",
        "page": page,
        "summary": summary,
        "outline": outline[:200],
        "paragraphs": paragraphs,
        "truncated": len(doc.paragraphs) > MAX_PARA,
    }


# ── pdf (PyMuPDF / fitz, optional) ───────────────────────────────────────────

def _inspect_pdf(path):
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return {
            "success": False,
            "needsDep": True,
            "error": "PDF 格式提取需要 PyMuPDF。",
            "hint": "安装：pip install khy-os[doc]  或  pip install pymupdf。安装后重试。",
        }

    doc = fitz.open(path)
    pages = []
    font_counter = {}
    size_counter = {}
    MAX_PAGES = 50
    for pno in range(min(doc.page_count, MAX_PAGES)):
        page = doc.load_page(pno)
        try:
            data = page.get_text("dict")
        except Exception:
            continue
        spans_sample = []
        for block in data.get("blocks", []):
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    fnt = span.get("font")
                    sz = round(float(span.get("size", 0)), 1) if span.get("size") else None
                    if fnt:
                        font_counter[fnt] = font_counter.get(fnt, 0) + 1
                    if sz:
                        size_counter[sz] = size_counter.get(sz, 0) + 1
                    if len(spans_sample) < 5 and span.get("text", "").strip():
                        spans_sample.append({"text": span.get("text", "")[:60], "font": fnt, "sizePt": sz})
        pages.append({"page": pno + 1, "spansSample": spans_sample})

    def _dominant(counter):
        return max(counter.items(), key=lambda kv: kv[1])[0] if counter else None

    summary = {
        "pageCount": doc.page_count,
        "dominantFont": _dominant(font_counter),
        "dominantSizePt": _dominant(size_counter),
    }
    doc.close()
    return {"success": True, "format": "pdf", "summary": summary, "pages": pages,
            "truncated": doc.page_count > MAX_PAGES if hasattr(doc, "page_count") else False}


# ── plain text / markdown / source ───────────────────────────────────────────

def _inspect_text(path, ext):
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()
    except Exception as e:
        return {"success": False, "error": "read failed: %s" % e}

    lines = content.splitlines()
    fmt = "markdown" if ext in (".md", ".markdown") else "text"
    outline = []
    if fmt == "markdown":
        for ln in lines:
            s = ln.lstrip()
            if s.startswith("#"):
                level = len(s) - len(s.lstrip("#"))
                title = s[level:].strip()
                if 1 <= level <= 6 and title:
                    outline.append({"level": level, "text": title[:120]})

    paragraphs = [p for p in content.split("\n\n") if p.strip()]
    return {
        "success": True,
        "format": fmt,
        "note": "纯文本/源码无字体字号等视觉格式属性；以下为结构信息。",
        "summary": {
            "lineCount": len(lines),
            "paragraphCount": len(paragraphs),
            "charCount": len(content),
            "headingCount": len(outline),
        },
        "outline": outline[:200],
    }


def inspect(path):
    if not os.path.exists(path):
        _err("File not found: %s" % path)
        return
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".docx":
            try:
                from docx import Document  # noqa: F401
            except ImportError:
                _err("python-docx not available. Run: pip install khy-os[doc]", needsDep=True)
                return
            _output(_inspect_docx(path))
        elif ext == ".pdf":
            _output(_inspect_pdf(path))
        else:
            _output(_inspect_text(path, ext))
    except Exception as e:
        _err("Inspection failed: %s" % e, format=ext.lstrip("."))


if __name__ == "__main__":
    if len(sys.argv) < 3 or sys.argv[1] != "inspect":
        _err("Usage: docInspect.py inspect <path>")
        sys.exit(1)
    inspect(sys.argv[2])
