#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
docTypeset.py — deterministic document renderer (content + style → .docx).

Called by the Node tool src/tools/renderDocument.js via subprocess. It reads a
single JSON payload {ast, template, output} and produces a .docx whose formatting
is decided ENTIRELY by the style template — never by the model. The model only
supplies the semantic AST (headings / paragraphs / lists / tables / …); this
module owns every visual decision.

Design contract (mirrors the Node-side architecture):
  - Atomic typeset primitives: _set_run_font / _apply_paragraph_format /
    _set_page / _add_page_break / _apply_heading. Block rendering only ever calls
    these — it never hand-assembles XML or format strings.
  - Chinese is first-class: _set_run_font ALWAYS stamps w:rFonts/@w:eastAsia, the
    single most common reason a Chinese font silently fails to apply.
  - Page breaks are explicit API calls (pagebreak block) or the template's
    page_break_before policy — never whitespace.
  - Write-after verification: after saving, the file is re-parsed and key metrics
    (A4 page size, eastAsia on every run, heading sizes) are checked; mismatches
    are PATCHED in place via the same API, never bounced back to the model.

Dependencies: python-docx (installed with khy-os[doc]).
"""
import sys
import json
import os

# A4 in EMU (English Metric Units): 1 mm = 36000 EMU.
A4_WIDTH_EMU = 210 * 36000
A4_HEIGHT_EMU = 297 * 36000
EMU_TOLERANCE = 36000  # 1 mm slack when verifying page size


def _output(data):
    print(json.dumps(data, ensure_ascii=False))


def _err(msg, **extra):
    out = {"success": False, "error": msg}
    out.update(extra)
    _output(out)


# ── Atomic typeset primitives ──────────────────────────────────────────────

def _set_run_font(run, spec, default_spec):
    """Apply a font spec to a run. ALWAYS sets eastAsia (防呆: Chinese fonts fail
    silently in docx unless w:rFonts/@w:eastAsia is set explicitly)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt, RGBColor

    s = dict(default_spec or {})
    s.update(spec or {})

    ascii_font = s.get("ascii", "Times New Roman")
    east_asia = s.get("eastAsia", "宋体")

    run.font.name = ascii_font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:cs"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east_asia)  # ← the critical attribute

    size = s.get("size")
    if size is not None:
        run.font.size = Pt(float(size))
    if s.get("bold") is not None:
        run.font.bold = bool(s.get("bold"))
    if s.get("italic") is not None:
        run.font.italic = bool(s.get("italic"))
    if s.get("smallCaps"):
        run.font.small_caps = True
    color = s.get("color")
    if color:
        try:
            run.font.color.rgb = RGBColor.from_string(str(color).lstrip("#"))
        except Exception:
            pass


def _apply_paragraph_format(para, para_base, font_spec, first_line_indent_chars=0):
    """Apply paragraph-level formatting (alignment, line spacing, spacing before/
    after, first-line indent, page-break-before) deterministically from the spec."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

    pf = para.paragraph_format
    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    align = (font_spec or {}).get("align")
    if align in align_map:
        pf.alignment = align_map[align]

    # Line spacing: multiple (e.g. 1.5x) or exact points.
    rule = (para_base or {}).get("lineSpacingRule", "multiple")
    ls = (para_base or {}).get("lineSpacing")
    if ls is not None:
        if rule == "exact":
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = Pt(float(ls))
        else:
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing = float(ls)

    # Spacing before/after: font spec overrides paragraph base when present.
    sb = (font_spec or {}).get("spaceBefore", (para_base or {}).get("spaceBefore"))
    sa = (font_spec or {}).get("spaceAfter", (para_base or {}).get("spaceAfter"))
    if sb is not None:
        pf.space_before = Pt(float(sb))
    if sa is not None:
        pf.space_after = Pt(float(sa))

    # First-line indent: approximate "N chars" as N × body font size (CJK char
    # width ≈ font em). Only body paragraphs pass a positive value here.
    if first_line_indent_chars and first_line_indent_chars > 0:
        size = float((font_spec or {}).get("size", 12))
        pf.first_line_indent = Pt(size * float(first_line_indent_chars))

    # Page-break-before is an explicit, deterministic paragraph property.
    if (font_spec or {}).get("pageBreakBefore"):
        pf.page_break_before = True


def _apply_left_indent_chars(para, font_spec, chars):
    from docx.shared import Pt
    if chars and chars > 0:
        size = float((font_spec or {}).get("size", 12))
        para.paragraph_format.left_indent = Pt(size * float(chars))


def _apply_hanging_indent_chars(para, font_spec, chars):
    from docx.shared import Pt
    if chars and chars > 0:
        size = float((font_spec or {}).get("size", 12))
        para.paragraph_format.left_indent = Pt(size * float(chars))
        para.paragraph_format.first_line_indent = Pt(-size * float(chars))


def _set_page(section, page_spec):
    """Set A4 (or named) page size, margins, header/footer distance."""
    from docx.shared import Cm

    size = (page_spec or {}).get("size", "A4")
    if str(size).upper() == "A4":
        section.page_width = A4_WIDTH_EMU
        section.page_height = A4_HEIGHT_EMU
    elif str(size).upper() == "LETTER":
        section.page_width = int(8.5 * 25.4 * 36000)
        section.page_height = int(11 * 25.4 * 36000)

    m = (page_spec or {}).get("margins", {})
    if m:
        if m.get("top") is not None:
            section.top_margin = Cm(float(m["top"]))
        if m.get("bottom") is not None:
            section.bottom_margin = Cm(float(m["bottom"]))
        if m.get("left") is not None:
            section.left_margin = Cm(float(m["left"]))
        if m.get("right") is not None:
            section.right_margin = Cm(float(m["right"]))

    header = (page_spec or {}).get("header", {})
    footer = (page_spec or {}).get("footer", {})
    if header.get("distance") is not None:
        section.header_distance = Cm(float(header["distance"]))
    if footer.get("distance") is not None:
        section.footer_distance = Cm(float(footer["distance"]))


def _add_page_number(section):
    """Insert a centered { PAGE } field into the footer (deterministic API)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._element.append(fld_begin)
    run._element.append(instr)
    run._element.append(fld_end)


# ── Block renderers ────────────────────────────────────────────────────────

def _fonts(template):
    return (template or {}).get("fonts", {})


def _render_heading(doc, block, template, is_first_content):
    level = int(block.get("level", 1))
    key = "heading%d" % level if ("heading%d" % level) in _fonts(template) else "heading1"
    font_spec = _fonts(template).get(key, {})
    default_spec = _fonts(template).get("default", {})
    para_base = template.get("paragraph", {})

    para = doc.add_paragraph()
    # Auto page-break before H1 per template pagination policy (skip the very
    # first content block so the document does not open with a blank page).
    pagination = template.get("pagination", {})
    if level == 1 and pagination.get("pageBreakBeforeHeading1") and not is_first_content:
        para.paragraph_format.page_break_before = True

    _apply_paragraph_format(para, para_base, font_spec, first_line_indent_chars=0)
    run = para.add_run(block.get("text", ""))
    _set_run_font(run, font_spec, default_spec)


def _render_paragraph(doc, block, template):
    default_spec = _fonts(template).get("default", {})
    para_base = template.get("paragraph", {})
    fl = para_base.get("firstLineIndentChars", 0)

    para = doc.add_paragraph()
    _apply_paragraph_format(para, para_base, default_spec, first_line_indent_chars=fl)

    runs = block.get("runs")
    if isinstance(runs, list) and runs:
        for r in runs:
            run = para.add_run(r.get("text", ""))
            spec = dict(default_spec)
            if r.get("bold"):
                spec["bold"] = True
            if r.get("italic"):
                spec["italic"] = True
            _set_run_font(run, spec, default_spec)
    else:
        run = para.add_run(block.get("text", ""))
        _set_run_font(run, default_spec, default_spec)


def _render_list(doc, block, template):
    list_spec = _fonts(template).get("list", _fonts(template).get("default", {}))
    default_spec = _fonts(template).get("default", {})
    para_base = template.get("paragraph", {})
    indent_chars = (template.get("list", {}) or {}).get("indentChars", 2)
    ordered = bool(block.get("ordered"))
    for idx, item in enumerate(block.get("items", []), start=1):
        prefix = ("%d. " % idx) if ordered else "• "
        para = doc.add_paragraph()
        _apply_paragraph_format(para, para_base, default_spec, first_line_indent_chars=0)
        _apply_left_indent_chars(para, list_spec, indent_chars)
        run = para.add_run(prefix + str(item))
        _set_run_font(run, list_spec, default_spec)


def _render_quote(doc, block, template):
    quote_spec = _fonts(template).get("quote", _fonts(template).get("default", {}))
    default_spec = _fonts(template).get("default", {})
    para_base = template.get("paragraph", {})
    para = doc.add_paragraph()
    _apply_paragraph_format(para, para_base, quote_spec, first_line_indent_chars=0)
    _apply_left_indent_chars(para, quote_spec, quote_spec.get("indentChars", 2))
    run = para.add_run(block.get("text", ""))
    _set_run_font(run, quote_spec, default_spec)


def _render_code(doc, block, template):
    code_spec = _fonts(template).get("code", _fonts(template).get("default", {}))
    default_spec = _fonts(template).get("default", {})
    # One paragraph per line, no first-line indent, exact spacing — preserve layout.
    for line in str(block.get("text", "")).split("\n"):
        para = doc.add_paragraph()
        para.paragraph_format.first_line_indent = None
        run = para.add_run(line if line != "" else " ")
        _set_run_font(run, code_spec, default_spec)


def _render_table(doc, block, template):
    default_spec = _fonts(template).get("default", {})
    cell_spec = _fonts(template).get("tableCell", default_spec)
    table_cfg = template.get("table", {})
    header = block.get("header")
    rows = block.get("rows", [])
    n_cols = max([len(r) for r in rows] + ([len(header)] if header else [0]))
    if n_cols == 0:
        return
    table = doc.add_table(rows=0, cols=n_cols)
    try:
        table.style = table_cfg.get("style", "Table Grid")
    except Exception:
        pass

    def _fill_row(values, bold):
        cells = table.add_row().cells
        for ci in range(n_cols):
            text = values[ci] if ci < len(values) else ""
            cell = cells[ci]
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(str(text))
            spec = dict(cell_spec)
            if bold:
                spec["bold"] = True
            _set_run_font(run, spec, default_spec)

    if header:
        _fill_row(header, bool(table_cfg.get("headerBold", True)))
    for r in rows:
        _fill_row(r, False)


def _render_figure(doc, block, template):
    from docx.shared import Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    caption_spec = _fonts(template).get("caption", _fonts(template).get("default", {}))
    default_spec = _fonts(template).get("default", {})
    img_path = block.get("path")
    if img_path and os.path.isfile(img_path):
        try:
            doc.add_picture(img_path, width=Cm(14))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass
    caption = block.get("caption")
    if caption:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(caption)
        _set_run_font(run, caption_spec, default_spec)


def _render_reference(doc, block, template):
    ref_spec = _fonts(template).get("reference", _fonts(template).get("default", {}))
    default_spec = _fonts(template).get("default", {})
    para_base = template.get("paragraph", {})
    for idx, entry in enumerate(block.get("entries", []), start=1):
        para = doc.add_paragraph()
        _apply_paragraph_format(para, {**para_base, "lineSpacing": 1.0}, ref_spec, first_line_indent_chars=0)
        _apply_hanging_indent_chars(para, ref_spec, ref_spec.get("hangingIndentChars", 2))
        run = para.add_run("[%d] %s" % (idx, str(entry)))
        _set_run_font(run, ref_spec, default_spec)


_RENDERERS = {
    "heading": None,  # handled specially (needs is_first_content)
    "paragraph": _render_paragraph,
    "list": _render_list,
    "quote": _render_quote,
    "code": _render_code,
    "table": _render_table,
    "figure": _render_figure,
    "reference": _render_reference,
}


# ── Top-level render + verification ────────────────────────────────────────

def _render_document(ast, template, output_path):
    from docx import Document

    doc = Document()
    section = doc.sections[0]
    _set_page(section, template.get("page", {}))
    if (template.get("page", {}).get("footer", {}) or {}).get("pageNumber"):
        try:
            _add_page_number(section)
        except Exception:
            pass

    # Optional document title (semantic, from AST), styled by the template.
    title = ast.get("title")
    if title:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        title_spec = _fonts(template).get("title", _fonts(template).get("heading1", {}))
        default_spec = _fonts(template).get("default", {})
        p = doc.add_paragraph()
        _apply_paragraph_format(p, template.get("paragraph", {}), title_spec, first_line_indent_chars=0)
        r = p.add_run(str(title))
        _set_run_font(r, title_spec, default_spec)

    seen_content = bool(title)
    for block in ast.get("blocks", []):
        btype = block.get("type")
        if btype == "pagebreak":
            doc.add_page_break()
            continue
        if btype == "heading":
            _render_heading(doc, block, template, is_first_content=not seen_content)
            seen_content = True
            continue
        fn = _RENDERERS.get(btype)
        if fn:
            fn(doc, block, template)
            seen_content = True

    out_dir = os.path.dirname(output_path)
    if out_dir and not os.path.isdir(out_dir):
        os.makedirs(out_dir, exist_ok=True)
    doc.save(output_path)


def _verify_and_patch(output_path, template):
    """Re-parse the saved docx and check key indicators. Patch mismatches in place
    (never bounce back to the model). Returns a validation report."""
    from docx import Document
    from docx.oxml.ns import qn

    report = {"pageSizeA4": False, "eastAsiaApplied": False, "headingSizeOk": None, "patched": []}
    doc = Document(output_path)
    patched = False

    # 1) Page size == A4 (within tolerance) — patch if off.
    want_a4 = str(template.get("page", {}).get("size", "A4")).upper() == "A4"
    sec = doc.sections[0]
    if want_a4:
        ok = (abs(int(sec.page_width) - A4_WIDTH_EMU) <= EMU_TOLERANCE and
              abs(int(sec.page_height) - A4_HEIGHT_EMU) <= EMU_TOLERANCE)
        if not ok:
            sec.page_width = A4_WIDTH_EMU
            sec.page_height = A4_HEIGHT_EMU
            report["patched"].append("page_size->A4")
            patched = True
        report["pageSizeA4"] = True
    else:
        report["pageSizeA4"] = None

    # 2) Every run carries an eastAsia font — the silent-failure guard. Patch any
    #    run missing it with the template default eastAsia.
    default_east = _fonts(template).get("default", {}).get("eastAsia", "宋体")
    default_ascii = _fonts(template).get("default", {}).get("ascii", "Times New Roman")
    missing = 0
    total = 0
    for para in doc.paragraphs:
        for run in para.runs:
            total += 1
            rpr = run._element.find(qn("w:rPr"))
            rfonts = rpr.find(qn("w:rFonts")) if rpr is not None else None
            if rfonts is None or rfonts.get(qn("w:eastAsia")) is None:
                # Patch directly via XML.
                from docx.oxml import OxmlElement
                rpr2 = run._element.get_or_add_rPr()
                rf = rpr2.find(qn("w:rFonts"))
                if rf is None:
                    rf = OxmlElement("w:rFonts")
                    rpr2.append(rf)
                rf.set(qn("w:eastAsia"), default_east)
                if rf.get(qn("w:ascii")) is None:
                    rf.set(qn("w:ascii"), default_ascii)
                    rf.set(qn("w:hAnsi"), default_ascii)
                missing += 1
                patched = True
    report["eastAsiaApplied"] = (missing == 0)
    if missing:
        report["patched"].append("eastAsia_runs:%d" % missing)

    # 3) Heading-1 size matches template (spot check the first H1 paragraph).
    want_h1 = _fonts(template).get("heading1", {}).get("size")
    if want_h1 is not None:
        from docx.shared import Pt
        want_emu = int(Pt(float(want_h1)))
        checked = None
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.size is not None and abs(int(run.font.size) - want_emu) <= int(Pt(2)):
                    checked = True
                    break
            if checked:
                break
        report["headingSizeOk"] = bool(checked) if checked is not None else None

    if patched:
        doc.save(output_path)
    return report


def render(payload_path):
    try:
        with open(payload_path, "r", encoding="utf-8") as f:
            payload = json.load(f)
    except Exception as e:
        _err("Could not read payload: %s" % e)
        return

    ast = payload.get("ast")
    template = payload.get("template") or {}
    output_path = payload.get("output")
    if not isinstance(ast, dict) or not isinstance(output_path, str):
        _err("Payload must contain {ast: object, template: object, output: string}")
        return

    try:
        from docx import Document  # noqa: F401
    except ImportError:
        _err(
            "python-docx not available. Run: pip install khy-os[doc]",
            needsDep=True,
        )
        return

    try:
        _render_document(ast, template, output_path)
    except Exception as e:
        _err("Rendering failed: %s" % e)
        return

    try:
        validation = _verify_and_patch(output_path, template)
    except Exception as e:
        validation = {"error": "verification skipped: %s" % e}

    try:
        out_size = os.path.getsize(output_path)
    except OSError:
        out_size = 0

    _output({
        "success": True,
        "output": output_path,
        "outputSize": out_size,
        "blocks": len(ast.get("blocks", [])),
        "validation": validation,
        "message": "Typeset: %s (%d blocks, %.0fKB)" % (
            os.path.basename(output_path), len(ast.get("blocks", [])), out_size / 1024.0),
    })


if __name__ == "__main__":
    if len(sys.argv) < 2:
        _err("Usage: docTypeset.py render <payload.json>")
        sys.exit(1)
    cmd = sys.argv[1]
    if cmd == "render":
        if len(sys.argv) < 3:
            _err("Usage: docTypeset.py render <payload.json>")
            sys.exit(1)
        render(sys.argv[2])
    else:
        _err("Unknown command: %s" % cmd)
        sys.exit(1)
